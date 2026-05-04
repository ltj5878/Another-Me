from pathlib import Path
import re
from uuid import UUID

from fastapi import HTTPException, UploadFile, status
from docx import Document
from sqlalchemy.orm import Session

from app.models.article import ArticleChunk, SourceArticle
from app.db.session import SessionLocal
from app.models.style import StyleCategory
from app.services.chunking import chunk_article
from app.services.embeddings import count_tokens, get_embeddings

SUPPORTED_EXTENSIONS = {".txt", ".md", ".docx"}
SOURCE_TYPES = {
    ".txt": "txt",
    ".md": "markdown",
    ".docx": "docx",
}

FULLWIDTH_TO_HALFWIDTH = str.maketrans(
    "！（）：；［］｛｝",
    "!():;[]{}",
)

_CONTROL_CHAR_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]")
_REPLACEMENT_CHAR_RE = re.compile(r"[�￾￿]")
_PAGE_NUMBER_RE = re.compile(r"^\s*[-—]\s*\d+\s*[-—]\s*$", re.MULTILINE)
_PAGE_CN_RE = re.compile(r"^\s*第\s*\d+\s*页.*$", re.MULTILINE)
_HEADER_FOOTER_RE = re.compile(
    r"^\s*(来源|编辑|责编|记者|转载|原文链接|版权所有)[:：].*$",
    re.MULTILINE,
)
_COPYRIGHT_CN_RE = re.compile(r"©.*?版权所有[^\n]*")
_COPYRIGHT_EN_RE = re.compile(r"Copyright\s*©?.*All\s*Rights\s*Reserved[^\n]*", re.IGNORECASE)
_COPYRIGHT_ARTICLE_RE = re.compile(r"本文(版权|著作权)归.*所有[^\n]*")
_BLANK_LINES_RE = re.compile(r"\n{3,}")


def clean_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = _CONTROL_CHAR_RE.sub("", text)
    text = _REPLACEMENT_CHAR_RE.sub("", text)
    text = text.translate(FULLWIDTH_TO_HALFWIDTH)
    text = _PAGE_NUMBER_RE.sub("", text)
    text = _PAGE_CN_RE.sub("", text)
    text = _HEADER_FOOTER_RE.sub("", text)
    text = _COPYRIGHT_CN_RE.sub("", text)
    text = _COPYRIGHT_EN_RE.sub("", text)
    text = _COPYRIGHT_ARTICLE_RE.sub("", text)
    lines = [line.strip() for line in text.split("\n")]
    text = "\n".join(lines)
    text = _BLANK_LINES_RE.sub("\n\n", text)
    return text.strip()


def count_words(text: str) -> int:
    chinese_chars = re.findall(r"[一-鿿]", text)
    latin_words = re.findall(r"[A-Za-z0-9]+(?:[-'][A-Za-z0-9]+)?", text)
    return len(chinese_chars) + len(latin_words)


def extract_text(file: UploadFile, suffix: str) -> str:
    raw = file.file.read()
    if suffix in {".txt", ".md"}:
        try:
            return raw.decode("utf-8")
        except UnicodeDecodeError as exc:
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="File must be valid UTF-8 text") from exc

    if suffix == ".docx":
        try:
            from io import BytesIO

            document = Document(BytesIO(raw))
        except Exception as exc:
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Unable to read .docx file") from exc

        text_parts = []
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
                
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():
                            text_parts.append(paragraph.text)
                            
        return "\n\n".join(text_parts)

    raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Unsupported file type")


def upload_source_article(db: Session, style_id: UUID, file: UploadFile) -> SourceArticle:
    style = db.get(StyleCategory, style_id)
    if style is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Style category not found")

    filename = file.filename or "untitled.txt"
    suffix = Path(filename).suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Only .txt, .md, and .docx files are supported")

    raw_text = extract_text(file, suffix)
    if not raw_text.strip():
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="File content is empty")

    article = SourceArticle(
        style_category_id=style.id,
        title=Path(filename).stem[:255],
        source_type=SOURCE_TYPES[suffix],
        original_filename=filename[:255],
        raw_text=raw_text.strip(),
        cleaned_text="",
        word_count=0,
        status="uploaded",
    )
    db.add(article)
    db.commit()
    db.refresh(article)
    return article


def process_source_article(article_id: UUID) -> None:
    with SessionLocal() as db:
        article = db.get(SourceArticle, article_id)
        if article is None:
            return
        style = db.get(StyleCategory, article.style_category_id)
        if style is None:
            _mark_article_failed(db, article_id)
            return

        try:
            article.status = "cleaning"
            db.commit()

            cleaned = clean_text(article.raw_text)
            if not cleaned:
                _mark_article_failed(db, article_id)
                return
            article.cleaned_text = cleaned
            article.word_count = count_words(cleaned)
            article.status = "chunking"
            db.commit()

            db.query(ArticleChunk).filter(ArticleChunk.source_article_id == article.id).delete()
            chunk_results = chunk_article(cleaned, article.title)

            article.status = "embedding"
            db.commit()

            texts = [c.content for c in chunk_results]
            embeddings = get_embeddings(texts) if texts else []

            for cr, emb in zip(chunk_results, embeddings):
                chunk = ArticleChunk(
                    source_article_id=article.id,
                    style_category_id=style.id,
                    chunk_index=cr.chunk_index,
                    content=cr.content,
                    embedding=emb,
                    token_count=count_tokens(cr.content),
                    chunk_metadata={
                        "article_title": article.title,
                        "position": cr.position,
                        "paragraph_count": cr.paragraph_count,
                        "writing_type": style.writing_type_hint,
                    },
                )
                db.add(chunk)

            article.status = "completed"
            db.commit()
        except Exception:
            db.rollback()
            _mark_article_failed(db, article_id)


def _mark_article_failed(db: Session, article_id: UUID) -> None:
    article = db.get(SourceArticle, article_id)
    if article is None:
        return
    article.status = "failed"
    db.commit()
